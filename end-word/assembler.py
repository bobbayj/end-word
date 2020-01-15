# Standard imports
import os, shutil
import xml.etree.ElementTree as ET

# Third-party imports
import docx  # To read docx and extract data
from docxcompose.composer import Composer  # Append files together, preserving everything except sections
from docxtpl import DocxTemplate, InlineImage

class Assembler:
    def __init__(self, dest, context, backpage, output_path):
        self.dest = dest
        self.context = context
        self.backpage = backpage
        self.output_path = output_path
    # def docx_composer(base, new_docx, new_page=False):
    #     '''Appends a new docx file to a base DocxTemplate object, and returns the object for further
        
    #     Supports text, styles, shapes, in-line images and floating images
    #     DOES NOT support section breaks and columns
        
    #     Parameters
    #     ----------
    #     base : DocxTemplate object
    #         Base docx file being worked on
    #     new_docx : str
    #         The file location of the target docx source
        
    #     Returns
    #     -------
    #     combined_dest : DocxTemplate object
        
    #     '''
    #     # New page if true
    #     if new_page:
    #         base.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        
    #     composer = Composer(base)
    #     new_doc = docx.Document(new_docx)
    #     composer.append(new_doc)
        
    #     temp_file_path = os.path.join(temp_path,'combined.docx')
    #     composer.save(temp_file_path)
    #     combined_base = DocxTemplate(temp_file_path)
        
    #     return combined_base


    def publish(self):
        
        # Finalise destination
        dest.paragraphs[-1].add_run().add_break(WD_BREAK.PAGE)
        dest.render(context)
        
        # Finalise backpage
        backpage_doc = DocxTemplate(backpage)
        backpage_doc.render(context)
        
        # Combine documents
        composer = Composer(dest)
        composer.append(backpage_doc)
        
        # Save output and delete temp folder with all contents
        composer.save(output_path)
        print(f'Saved at {output_path}')
        shutil.rmtree(temp_path)


    def append_xlsx(self, dest, source, heading=None):
        '''Appends Excel data source to the destination Word doc as a Table
        
        Does not dynamically search for table contents.
        Table data must start in cell A1.
        Only plots one table per Excel workbook.
        
        Parameters
        ----------
        dest : str
            The file location of the destination word doc
        source: str
            The file location of the target Excel source
        heading: str
            A string that will be printed in the style of Heading 1 above the table in word (default is None)
        '''
        # Sub-functions
        def xl2doc_color(color_meta):
            if color_meta is None:
                return
            tint = color_meta.tint
            if color_meta.type == 'theme':
                theme = color_meta.theme
                fillcolor = theme_and_tint_to_rgb(wb, theme, tint)
            elif color_meta.type == 'rgb':
                ms_rgb = color_meta.rgb
                fillcolor = ms_rgb_to_hex_rgb(ms_rgb, tint)
            elif color_meta.type == 'indexed':
                index = color_meta.indexed
                ms_rgb = styles.colors.COLOR_INDEX[index]
                if 'Foreground' in ms_rgb:
                    ms_rgb = '00000000'  # Black
                else:
                    ms_rgb = 'FF000000' # White
                fillcolor = ms_rgb_to_hex_rgb(ms_rgb, tint)
                
            else:
                raise TypeError(f'Unrecognised color-type: "{color_meta.type}". Check classes')
            return fillcolor

        
        def cell_text_runs():
            '''Copies cell text runs (for subscript and superscript support)
            '''
            wb = custom_load_workbook(source) # Related to Excel XML parser module
            ws = wb.active
            max_col = ws.dim['col_last']
            max_row = ws.dim['rw_last']
            
            # Store cell contents in an array
            cell_contents = []
            
            ws_range = CellHelpers.build_range(1, max_row,1,max_col)
            for r_idx, row in enumerate(ws.get_range(ws_range)):
                row_contents = []
                for c_idx, cell in enumerate(row):
                    # Convert cell reference to SharedString pointer
                    cell_ref = ws.cell(cell.rw, cell.col)

                    # Now we can store the SharedString value
                    cell_value = cell_ref.value
                    row_contents.append(cell_value)
                cell_contents.append(row_contents)
            
            # Write cell contents to the docx using add_to_paragraph()
    #         cell_value.add_to_paragraph(paragraph)
            
            return cell_contents
        
        
        # Openpyxl
        # Note: openpyxl cannot read/copy charts; it needs to recreate them from source data
        # Read-only and data-only increases the speed of reading data from workbooks
        wb = load_workbook(source, data_only=True)
        
        table_values = cell_text_runs()
        
        # Loop through each worksheet
        for ws in wb.worksheets:

            # Get dimensions of table in Excel
            data_vals = np.asarray(tuple(ws.values))  # Will need to add logic to read only the table and not comments
            table_dim = np.shape(data_vals)

            # Get merged ranges
            merged_ranges = ws.merged_cells.ranges

            # Store dict of formats
            src_fmts = {}
            for r,row in enumerate(ws.rows):
                for c,cell in enumerate(row):
                    src_fmts[(r,c)] = {
                        'bold': cell.font.b,
                        'italic': cell.font.i,
                        'name': cell.font.name,
                        'size': cell.font.size,
                        'fillColor': xl2doc_color(cell.fill.start_color),
                        'fontColor': xl2doc_color(cell.font.color),
                        'horizontal': cell.alignment.horizontal,
                        'vertical': cell.alignment.vertical,  # can build overrides
                        'border': {
                            'top': cell.border.top.style,
                            'topColor': xl2doc_color(cell.border.top.color),
                            'bottom': cell.border.bottom.style,
                            'bottomColor': xl2doc_color(cell.border.bottom.color),
                            'left': cell.border.left.style,
                            'leftColor': xl2doc_color(cell.border.left.color),
                            'right': cell.border.right.style,
                            'rightColor': xl2doc_color(cell.border.right.color),
                        }
                    }
            # Docx
            new_section_cols(dest, 1)  # Ensure Word section has only one column

            # Add heading if required
            if heading:
                dest.add_paragraph(style='Heading 1').add_run().add_text(heading)

            # Create, table in word
            table = dest.add_table(rows=table_dim[0], cols=table_dim[1])
            
            # Merge table cells if any found in Excel
            if len(merged_ranges):
                for merged_range in merged_ranges:
                    start_cell = table.cell(
                        merged_range.min_row-1,
                        merged_range.min_col-1
                    )
                    end_cell = table.cell(
                        merged_range.max_row-1,
                        merged_range.max_col-1
                    )
                    start_cell.merge(end_cell)
                    
            # Write to table
            for r, row in enumerate(table.rows):
                for c, cell in enumerate(row.cells):
                    if len(table_values[r][c].plain_text()) > 0:
                        table_values[r][c].add_to_paragraph(cell.paragraphs[0])
                        
            # Style table
            style_tbl(table, src_fmts)

    def append_docx(self, dest, data, columns=1, new_page=False, separate_header=False):
        '''Appends content from the Word source to the destination Word doc - supports text and in-line images.
        DOES NOT SUPPORT FLOATING IMAGES AND SHAPES! Use add_docx() instead
        
        Parameters
        ----------
        dest : str
            The file location of the destination word doc
        source: str
            The file location of the target Word source
        '''
        source = docx.Document(data)
        ims = [im for im in source.inline_shapes]

        # Persistent indexes
        im_paths = []
        im_heights = []
        im_widths = []

        # Temp variables
        im_streams = []
        
        # New page if true
        if new_page:
            dest.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        
        # Get image binary and metadata
        for im_idx, im in enumerate(ims):
            # Binary
            blip = im._inline.graphic.graphicData.pic.blipFill.blip
            rId = blip.embed
            doc_part = source.part
            image_part = doc_part.related_parts[rId]
            byte_data = image_part._blob
            image_stream = io.BytesIO(byte_data)
            im_streams.append(image_stream)

            # Metadata
            image_name = f'img_{im_idx}.jpeg'
            im_heights.append(im.height.mm)
            im_widths.append(im.width.mm)
            
            # Save images to temp folder
            im_path = os.path.join(temp_path,image_name)
            im_paths.append(im_path)
            with open(im_path, "wb") as fh:
                fh.write(byte_data)
            fh.close()
        # ----End image transfer----
        # ----Start copying docx content----
        # Populate and save output
        paras = source.paragraphs
        im_idx = 0
        
        # Split into columns if header is not separate
        # Otherwise, split into columns after the header
        if not separate_header:
            new_section_cols(dest, columns)
        
        for para_idx, para in enumerate(paras):
            if(para.text):
                get_para_data(dest, para)
            
            # Copy images over
            root = ET.fromstring(para._p.xml)
            namespace = {'wp':"http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"}
            inlines = root.findall('.//wp:inline',namespace)

            if(len(inlines) > 0):
                uid = f'img_{im_idx}'

                img = dest.add_paragraph()
                img.add_run().add_text("{{ " + uid + " }}")
                img.alignment = WD_ALIGN_PARAGRAPH.CENTER

                context[uid] = InlineImage(
                    dest,
                    im_paths[im_idx],
                    width=Mm(im_widths[im_idx]),
                    height=Mm(im_heights[im_idx]),
                )
                im_idx += 1
                
                
            # Split into columns after the header
            if (para_idx == 0 and separate_header):
                new_section_cols(dest, columns)
        

